#!/usr/bin/env python3
"""Apply ONKI manuscript page and paragraph settings to a Pandoc DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


BODY_FONT = "Yu Mincho"
LATIN_FONT = "Times New Roman"
BACK_MATTER_SECTION_TITLES = (
    "Table",
    "Tables",
    "Figure Legend",
    "Figure Legends",
    "Figure",
    "Figures",
)
TABLE1_HEADERS = ("Characteristic", "Category", "n (%)")
TABLE1_COLUMN_WIDTHS_MM = (45, 75, 30)
TABLE1_TWO_PANEL_COLUMN_WIDTHS_MM = (22, 36, 17, 22, 36, 17)
TABLE1_FOOTNOTE_PREFIX = "Values are presented as n (%)"
ENGLISH_NUMERIC_RANGE = re.compile(r"(?<=\d)-(?=\d)")
ENGLISH_SUMMARY_FIELDS = (
    "Authors: ",
    "Affiliations: ",
    "Corresponding author: ",
)


def apply_english_range_typography(text: str) -> str:
    """Use en dashes for ranges in rendered English submission content."""
    text = text.replace("January-February", "January–February")
    return ENGLISH_NUMERIC_RANGE.sub("–", text)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Format a Pandoc-generated manuscript DOCX as A4 with approximately "
            "35 Japanese characters per line and 32 lines per page."
        )
    )
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("--output-docx", type=Path)
    return parser.parse_args()


def set_run_font(run, size: Pt) -> None:
    run.font.name = LATIN_FONT
    run.font.size = size
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)


def set_style_font(style, size: Pt) -> None:
    style.font.name = LATIN_FONT
    style.font.size = size
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)


def mark_header_row(row) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table_row_properties.append(repeat)


def prevent_row_split(row) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    table_row_properties.append(cannot_split)


def get_or_add_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def clear_list_formatting(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    numbering = properties.find(qn("w:numPr"))
    if numbering is not None:
        properties.remove(numbering)


def set_paragraph_text(paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def move_after(paragraph, anchor) -> None:
    anchor._p.addnext(paragraph._p)


def add_page_number(section) -> None:
    section.footer_distance = Mm(12.7)
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, result, end):
        run._r.append(element)
    set_run_font(run, Pt(10.5))


def add_line_numbering(section) -> None:
    line_numbering = get_or_add_child(section._sectPr, "w:lnNumType")
    line_numbering.set(qn("w:countBy"), "1")


def prepare_english_summary(document) -> str | None:
    paragraphs = document.paragraphs
    heading = next(
        (paragraph for paragraph in paragraphs if paragraph.text.strip() == "English Title"),
        None,
    )
    if heading is None:
        return None

    index = paragraphs.index(heading)
    if index + 1 >= len(paragraphs):
        return None
    title_paragraph = paragraphs[index + 1]
    title = title_paragraph.text.strip()
    set_paragraph_text(heading, title)
    remove_paragraph(title_paragraph)

    paragraphs = document.paragraphs
    authors = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith(ENGLISH_SUMMARY_FIELDS[0])
    )
    affiliations = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith(ENGLISH_SUMMARY_FIELDS[1])
    )
    corresponding = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith(ENGLISH_SUMMARY_FIELDS[2])
    )
    abstract_heading = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip() == "English Abstract"
    )
    keywords = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith("Keywords:")
    )

    for paragraph, prefix in (
        (authors, ENGLISH_SUMMARY_FIELDS[0]),
        (affiliations, ENGLISH_SUMMARY_FIELDS[1]),
        (corresponding, ENGLISH_SUMMARY_FIELDS[2]),
    ):
        clear_list_formatting(paragraph)
        set_paragraph_text(paragraph, paragraph.text.strip().removeprefix(prefix))

    set_paragraph_text(abstract_heading, "Abstract")
    move_after(affiliations, keywords)
    move_after(corresponding, affiliations)

    return title


def set_table_fixed_width(table, widths_mm: tuple[int, ...]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_properties = table._tbl.tblPr
    layout = get_or_add_child(table_properties, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_width = get_or_add_child(table_properties, "w:tblW")
    table_width.set(qn("w:w"), str(Mm(sum(widths_mm)).twips))
    table_width.set(qn("w:type"), "dxa")

    for column, width_mm in zip(table.columns, widths_mm):
        width = Mm(width_mm)
        column.width = width
        for cell in column.cells:
            cell.width = width
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = get_or_add_child(cell_properties, "w:tcW")
            cell_width.set(qn("w:w"), str(width.twips))
            cell_width.set(qn("w:type"), "dxa")


def set_cell_margins(
    cell,
    *,
    top_mm: float,
    right_mm: float,
    bottom_mm: float,
    left_mm: float,
) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = get_or_add_child(cell_properties, "w:tcMar")
    for edge, value_mm in (
        ("top", top_mm),
        ("right", right_mm),
        ("bottom", bottom_mm),
        ("left", left_mm),
    ):
        margin = get_or_add_child(margins, f"w:{edge}")
        margin.set(qn("w:w"), str(Mm(value_mm).twips))
        margin.set(qn("w:type"), "dxa")


def set_table_borders_none(table) -> None:
    table_properties = table._tbl.tblPr
    borders = get_or_add_child(table_properties, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = get_or_add_child(borders, f"w:{edge}")
        border.set(qn("w:val"), "nil")


def set_cell_horizontal_borders(
    cell,
    *,
    top: tuple[str, str] | None = None,
    bottom: tuple[str, str] | None = None,
) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    borders = get_or_add_child(cell_properties, "w:tcBorders")
    for edge, specification in (("top", top), ("bottom", bottom)):
        border = get_or_add_child(borders, f"w:{edge}")
        if specification is None:
            border.set(qn("w:val"), "nil")
            continue
        size, color = specification
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def is_table1(table) -> bool:
    if not table.rows or len(table.rows[0].cells) != len(TABLE1_HEADERS):
        return False
    return tuple(cell.text.strip() for cell in table.rows[0].cells) == TABLE1_HEADERS


def is_two_panel_table1(table) -> bool:
    if not table.rows or len(table.rows[0].cells) != 2 * len(TABLE1_HEADERS):
        return False
    headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
    return headers == TABLE1_HEADERS + TABLE1_HEADERS


def format_table1(table) -> None:
    is_two_panel = is_two_panel_table1(table)
    column_widths = (
        TABLE1_TWO_PANEL_COLUMN_WIDTHS_MM
        if is_two_panel
        else TABLE1_COLUMN_WIDTHS_MM
    )
    set_table_fixed_width(table, column_widths)
    set_table_borders_none(table)

    panel_offsets = (0, len(TABLE1_HEADERS)) if is_two_panel else (0,)
    group_starts_by_panel = {
        panel_offset: [
            row_index
            for row_index in range(1, len(table.rows))
            if table.rows[row_index].cells[panel_offset].text.strip()
        ]
        for panel_offset in panel_offsets
    }

    for row_index, row in enumerate(table.rows):
        is_header = row_index == 0
        is_last_row = row_index == len(table.rows) - 1
        for column_index, cell in enumerate(row.cells):
            panel_offset = (
                len(TABLE1_HEADERS)
                if is_two_panel and column_index >= len(TABLE1_HEADERS)
                else 0
            )
            panel_column_index = column_index - panel_offset
            is_group_start = row_index in group_starts_by_panel[panel_offset]
            if is_two_panel:
                left_margin_mm = 1.5 if column_index == 3 else 0.5
                right_margin_mm = 1.5 if column_index == 2 else 0.5
            else:
                left_margin_mm = 1.5
                right_margin_mm = 1.5
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(
                cell,
                top_mm=0.5,
                right_mm=right_margin_mm,
                bottom_mm=0.5,
                left_mm=left_margin_mm,
            )
            set_cell_horizontal_borders(
                cell,
                top=("8", "000000") if is_header else (
                    ("4", "B7B7B7") if is_group_start else None
                ),
                bottom=("8", "000000") if is_header or is_last_row else None,
            )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.left_indent = (
                    None if is_two_panel else (
                        Mm(2.5)
                        if not is_header and panel_column_index == 1
                        else None
                    )
                )
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                    if panel_column_index == 2
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                if is_group_start and panel_column_index == 0:
                    for run in paragraph.runs:
                        run.font.bold = True

    for panel_offset in panel_offsets:
        group_starts = group_starts_by_panel[panel_offset]
        group_boundaries = group_starts + [len(table.rows)]
        for group_index, start_row in enumerate(group_starts):
            end_row = group_boundaries[group_index + 1]
            for row_index in range(start_row, end_row - 1):
                for column_index in range(
                    panel_offset,
                    panel_offset + len(TABLE1_HEADERS),
                ):
                    cell = table.rows[row_index].cells[column_index]
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True


def format_document(input_docx: Path, output_docx: Path) -> None:
    document = Document(input_docx)
    english_title = prepare_english_summary(document)

    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)
        add_page_number(section)
        add_line_numbering(section)

    styles = document.styles
    set_style_font(styles["Normal"], Pt(12))
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    styles["Normal"].paragraph_format.line_spacing = Pt(21.5)
    styles["Normal"].paragraph_format.space_after = Pt(0)

    heading_sizes = {
        "Title": Pt(15),
        "Heading 1": Pt(13),
        "Heading 2": Pt(12),
        "Heading 3": Pt(12),
    }
    for style_name, size in heading_sizes.items():
        if style_name in styles:
            set_style_font(styles[style_name], size)

    in_english_abstract = False
    in_english_section = False
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        is_table1_footnote = paragraph_text.startswith(TABLE1_FOOTNOTE_PREFIX)
        is_english_summary_title = bool(
            english_title and paragraph_text == english_title
        )
        is_english_summary_author = paragraph_text.startswith("Satoshi NOSHIRO")
        is_english_summary_affiliation = paragraph_text.startswith(
            "1) Aozora Clinic Kawagoe"
        )
        is_english_summary_corresponding = paragraph_text.startswith(
            "*Corresponding author"
        )
        is_english_summary_metadata = any(
            (
                is_english_summary_title,
                is_english_summary_author,
                is_english_summary_affiliation,
                is_english_summary_corresponding,
            )
        )
        if paragraph_text in ("和文抄録", "I はじめに"):
            paragraph.paragraph_format.page_break_before = True
        if is_english_summary_title:
            in_english_section = True
            paragraph.paragraph_format.page_break_before = True
        elif paragraph_text in BACK_MATTER_SECTION_TITLES:
            paragraph.paragraph_format.page_break_before = True
        if paragraph_text == "Abstract":
            in_english_abstract = True
        elif in_english_abstract and paragraph_text in BACK_MATTER_SECTION_TITLES:
            in_english_abstract = False

        if is_english_summary_title:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.keep_with_next = True
        elif is_english_summary_author:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.keep_with_next = True
        elif is_english_summary_affiliation or is_english_summary_corresponding:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True
        elif is_table1_footnote:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True
        elif has_drawing:
            drawing_heights = [
                int(extent.get("cy")) / 12700
                for extent in paragraph._p.xpath(".//wp:extent")
                if extent.get("cy")
            ]
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph.paragraph_format.line_spacing = Pt(max(drawing_heights))
        elif in_english_abstract and paragraph_text != "Abstract":
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            paragraph.paragraph_format.line_spacing = 2.0
        elif paragraph.style.name == "Normal":
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(21.5)

        for run in paragraph.runs:
            if in_english_section and not (
                is_english_summary_affiliation
                or is_english_summary_corresponding
            ):
                rendered_text = apply_english_range_typography(run.text)
                if rendered_text != run.text:
                    run.text = rendered_text
            if is_english_summary_affiliation or is_english_summary_corresponding:
                set_run_font(run, Pt(10))
            elif is_english_summary_metadata:
                set_run_font(run, Pt(12))
            elif is_table1_footnote:
                set_run_font(run, Pt(9))
            elif paragraph.style.name == "Title":
                set_run_font(run, Pt(15))
            elif paragraph.style.name.startswith("Heading"):
                set_run_font(run, Pt(12))
            else:
                set_run_font(run, Pt(12))

    for table in document.tables:
        if table.rows:
            mark_header_row(table.rows[0])
        for row in table.rows:
            prevent_row_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = (
                        WD_LINE_SPACING.SINGLE
                    )
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        rendered_text = apply_english_range_typography(run.text)
                        if rendered_text != run.text:
                            run.text = rendered_text
                        set_run_font(run, Pt(9))
        if is_table1(table) or is_two_panel_table1(table):
            format_table1(table)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main() -> None:
    args = parse_args()
    input_docx = args.input_docx.resolve()
    output_docx = (
        args.output_docx.resolve() if args.output_docx else input_docx
    )
    if not input_docx.exists():
        raise FileNotFoundError(input_docx)
    format_document(input_docx, output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
