from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
PACKAGE = (
    ROOT
    / "deliverables"
    / "06_submission"
    / "20260821_165125_onki_short_report_submission"
)

TITLE = (
    "青森県五所川原市における浴室暖房乾燥機の未設置・未使用理由と"
    "浴室の温冷感"
)
AUTHOR = "野城 聡志"
PRIMARY_AFFILIATION = "医療法人社団優青会 あおぞらクリニック川越"
SECONDARY_AFFILIATION = "一般社団法人 高齢者入浴アドバイザー協会"


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "BIZ UDPMincho"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDP明朝 Medium")
    run.font.size = Pt(size)
    run.font.bold = bold


def configure_a4(document: Document, top: float = 20, bottom: float = 20) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    normal = document.styles["Normal"]
    normal.font.name = "BIZ UDPMincho"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDP明朝 Medium")
    normal.font.size = Pt(10.5)


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    align=None,
    bold: bool = False,
    size: float = 10.5,
    space_before: float = 0,
    space_after: float = 5,
    line_spacing: float = 1.2,
):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def build_cover_letter() -> None:
    document = Document()
    configure_a4(document, top=18, bottom=18)

    add_paragraph(
        document,
        "投稿依頼文",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=15,
        space_after=10,
    )
    add_paragraph(
        document,
        "2026年　　月　　日",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=8,
    )
    add_paragraph(document, "日本温泉気候物理医学会雑誌", space_after=0)
    add_paragraph(document, "編集委員長　殿", space_after=8)
    add_paragraph(
        document,
        "短報投稿のお願い",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=12,
        space_after=8,
    )

    add_paragraph(
        document,
        "下記論文を短報として投稿いたします。本論文は未発表であり、他誌に投稿中または掲載予定のものではありません。著者は本論文の内容を確認し、日本温泉気候物理医学会雑誌への投稿に同意しています。",
    )
    add_paragraph(
        document,
        "本研究は、青森県五所川原市内で実施した無記名自記式質問紙による横断研究です。浴室暖房乾燥機を設置していない、または暖房として使用していない回答者が挙げた理由を冬季の浴室の寒さ体感別に記述し、セントラル暖房の使用状況も併せて示しました。寒冷地における浴室の寒さ対策を検討する際に、設備の必要性だけでなく、費用や住宅・設置上の制約を把握する必要性を示す基礎資料になると考えます。",
    )
    add_paragraph(
        document,
        "本研究は日本温泉気候物理医学会倫理審査会の承認を得て実施しました（承認番号Onki26-01、承認日2026年3月11日）。回答は任意かつ無記名とし、回答済み質問紙の提出をもって研究参加への同意としました。開示すべき利益相反はなく、外部研究費の提供も受けていません。英文題名、英文抄録、Table 1、Fig. 1 legendおよびFig. 1の英文はネイティブ校閲を受けています。",
    )
    add_paragraph(document, "ご査読のほど、よろしくお願い申し上げます。", space_after=8)

    add_paragraph(document, "記", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=4)
    add_paragraph(document, f"投稿種別：短報", space_after=2)
    add_paragraph(document, f"論文題目：{TITLE}", space_after=2)
    add_paragraph(document, f"筆頭著者：{AUTHOR}", space_after=2)
    add_paragraph(
        document,
        f"所属：1）{PRIMARY_AFFILIATION}　2）{SECONDARY_AFFILIATION}",
        space_after=2,
    )
    add_paragraph(
        document,
        "所在地：〒350-0045 埼玉県川越市南通町19-2 307",
        space_after=2,
    )
    add_paragraph(document, "電話：049-227-3414　Fax：049-227-3415", space_after=2)
    add_paragraph(document, "E-mail：nsatoshi.ebaa@gmail.com", space_after=6)
    add_paragraph(document, "署名：　　　　　　　　　　　　　　　　　", space_after=0)

    output = PACKAGE / "04_cover_letter_ready_to_sign.docx"
    document.save(output)


def add_checkbox_line(document: Document, text: str, checked: bool = False):
    mark = "☒" if checked else "☐"
    return add_paragraph(
        document,
        f"　{mark}　{text}",
        size=10,
        space_after=2,
        line_spacing=1.05,
    )


def build_ethics_checklist() -> None:
    document = Document()
    configure_a4(document, top=14, bottom=14)

    add_paragraph(document, f"論文題目：{TITLE}", size=10, space_after=2)
    add_paragraph(document, f"著者名：{AUTHOR}", size=10, space_after=6)
    add_paragraph(
        document,
        "倫理関係チェックリスト",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
        space_after=7,
    )

    add_paragraph(document, "1．ヒトを対象とする研究である", bold=True, size=10.5, space_after=2)
    add_paragraph(
        document,
        "（ヒトを対象とした介入研究、観察研究、アンケート研究、既存の診療情報・健康情報を用いる研究等を含む）",
        size=9,
        space_after=2,
        line_spacing=1.0,
    )
    add_checkbox_line(document, "該当する（設問2へ）", checked=True)
    add_checkbox_line(document, "該当しない（設問3へ）", checked=False)

    add_paragraph(document, "2．倫理委員会の審査を受けている", bold=True, size=10.5, space_before=2, space_after=2)
    add_checkbox_line(document, "所属機関内の審査を受けている", checked=False)
    add_checkbox_line(
        document,
        "日本温泉気候物理医学会倫理審査会の審査を受けている（承認番号：Onki26-01、承認日：2026年3月11日）",
        checked=True,
    )
    add_checkbox_line(document, "倫理審査承認の内容が論文内の方法の項に記載されている", checked=True)
    add_checkbox_line(
        document,
        "臨床試験の場合には、臨床試験登録公開システム（UMIN-CTR等）に登録されている（本研究は臨床試験に該当しない）",
        checked=False,
    )
    add_checkbox_line(document, "審査を受けていない", checked=False)
    add_paragraph(
        document,
        "※ ヒトを対象とした研究では、原則として倫理委員会などの審査を受け、被験者の同意を取得した上で実施されたものでなければならない（本誌投稿規定）。",
        size=8.5,
        space_after=4,
        line_spacing=1.0,
    )

    add_paragraph(document, "3．動物を対象とする研究である", bold=True, size=10.5, space_after=2)
    add_checkbox_line(document, "該当する（設問4へ）", checked=False)
    add_checkbox_line(document, "該当しない（チェック終了）", checked=True)

    add_paragraph(document, "4．専門委員会等の審査を受けている", bold=True, size=10.5, space_before=2, space_after=2)
    add_checkbox_line(document, "所属機関内の審査を受けている（承認番号：　　　　　　　　　）", checked=False)
    add_checkbox_line(document, "審査承認の内容が論文内の方法の項に記載されている", checked=False)
    add_checkbox_line(document, "審査を受けていない", checked=False)
    add_paragraph(
        document,
        "※ 動物を対象とする研究の場合に記入する項目。本研究は該当しない。",
        size=8.5,
        space_after=0,
        line_spacing=1.0,
    )

    output = PACKAGE / "06_onki_ethics_checklist_completed.docx"
    document.save(output)


def build_submission_agreement() -> None:
    source = PACKAGE / "05_onki_submission_agreement_unsigned.pdf"
    output = PACKAGE / "05_onki_submission_agreement_ready_to_sign.pdf"
    font_path = Path("C:/Windows/Fonts/BIZ-UDMinchoM.ttc")

    with tempfile.TemporaryDirectory(prefix="onki_agreement_") as temp_dir:
        prefix = Path(temp_dir) / "agreement"
        subprocess.run(
            [
                "pdftoppm",
                "-f",
                "1",
                "-singlefile",
                "-png",
                "-r",
                "300",
                str(source),
                str(prefix),
            ],
            check=True,
        )
        image = Image.open(prefix.with_suffix(".png")).convert("RGB")
        draw = ImageDraw.Draw(image)
        scale = 300 / 160
        font = ImageFont.truetype(str(font_path), int(18 * scale))
        small_font = ImageFont.truetype(str(font_path), int(17 * scale))

        def put(x: int, y: int, text: str, selected_font=font, spacing: int = 6):
            draw.multiline_text(
                (int(x * scale), int(y * scale)),
                text,
                font=selected_font,
                fill="black",
                spacing=int(spacing * scale),
            )

        put(260, 750, PRIMARY_AFFILIATION, small_font)
        put(
            260,
            1005,
            "青森県五所川原市における浴室暖房乾燥機の\n未設置・未使用理由と浴室の温冷感",
            small_font,
            spacing=5,
        )
        put(260, 1165, "なし（単著）", small_font)

        image.save(output, "PDF", resolution=300.0, quality=95)


def main() -> None:
    PACKAGE.mkdir(parents=True, exist_ok=True)
    build_cover_letter()
    build_submission_agreement()
    build_ethics_checklist()


if __name__ == "__main__":
    main()
